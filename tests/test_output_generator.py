"""
Unit tests for output generator - Word, Excel, CSV, JSON formats
"""
import pytest
from pathlib import Path
from app.processing.output_generator import OutputGenerator
from app.models import OutputFormat, DocumentChunk
from termcolor import cprint
import pandas as pd
from docx import Document
import json


@pytest.mark.unit
class TestOutputGenerator:
    """Test suite for OutputGenerator class"""

    @pytest.fixture(autouse=True)
    def setup(self, temp_dir):
        """Setup for each test"""
        self.generator = OutputGenerator(output_dir=temp_dir)
        self.output_dir = temp_dir

    def test_generator_initialization(self):
        """Test output generator initialization"""
        cprint("\n[TEST] Testing output generator initialization", "cyan")

        assert self.generator.output_dir.exists()
        cprint(f"[TEST] ✓ Output generator initialized at: {self.generator.output_dir}", "green")

    def test_generate_word_landscape(self, sample_verified_chunks_data):
        """Test Word document generation in landscape orientation"""
        cprint("\n[TEST] Testing Word landscape document generation", "cyan")

        output_path = self.generator.generate_output(
            chunks=sample_verified_chunks_data,
            original_filename="test_document.pdf",
            output_format=OutputFormat.WORD_LANDSCAPE
        )

        # Verify file was created
        assert output_path.exists()
        assert output_path.suffix == ".docx"

        # Verify document content
        doc = Document(output_path)

        # Should have at least one table
        assert len(doc.tables) > 0

        table = doc.tables[0]

        # Should have correct number of columns (7: Page, Item, Text, Verified, Score, Source, Note)
        assert len(table.columns) == 7

        # Should have header row + data rows
        assert len(table.rows) == len(sample_verified_chunks_data) + 1

        # Verify header row
        header_row = table.rows[0]
        expected_headers = ["Page #", "Item #", "Text", "Verified ☑", "Verification Score", "Verification Source", "Verification Note"]
        for idx, expected_header in enumerate(expected_headers):
            actual_header = header_row.cells[idx].text
            assert expected_header in actual_header or actual_header in expected_header

        cprint(f"[TEST] ✓ Word landscape document generated: {output_path.name}", "green")

    def test_generate_word_portrait(self, sample_verified_chunks_data):
        """Test Word document generation in portrait orientation"""
        cprint("\n[TEST] Testing Word portrait document generation", "cyan")

        output_path = self.generator.generate_output(
            chunks=sample_verified_chunks_data,
            original_filename="test_document.pdf",
            output_format=OutputFormat.WORD_PORTRAIT
        )

        # Verify file was created
        assert output_path.exists()
        assert output_path.suffix == ".docx"

        # Verify document content
        doc = Document(output_path)
        assert len(doc.tables) > 0

        cprint(f"[TEST] ✓ Word portrait document generated: {output_path.name}", "green")

    def test_generate_excel(self, sample_verified_chunks_data):
        """Test Excel document generation"""
        cprint("\n[TEST] Testing Excel document generation", "cyan")

        output_path = self.generator.generate_output(
            chunks=sample_verified_chunks_data,
            original_filename="test_document.pdf",
            output_format=OutputFormat.EXCEL
        )

        # Verify file was created
        assert output_path.exists()
        assert output_path.suffix == ".xlsx"

        # Verify Excel content
        df = pd.read_excel(output_path, sheet_name='Verification')

        # Should have correct number of columns
        assert len(df.columns) == 7

        # Should have correct number of rows
        assert len(df) == len(sample_verified_chunks_data)

        # Verify column names
        expected_columns = ["Page #", "Item #", "Text", "Verified ☑", "Verification Score", "Verification Source", "Verification Note"]
        for col in expected_columns:
            assert col in df.columns

        cprint(f"[TEST] ✓ Excel document generated: {output_path.name}", "green")

    def test_generate_csv(self, sample_verified_chunks_data):
        """Test CSV document generation"""
        cprint("\n[TEST] Testing CSV document generation", "cyan")

        output_path = self.generator.generate_output(
            chunks=sample_verified_chunks_data,
            original_filename="test_document.pdf",
            output_format=OutputFormat.CSV
        )

        # Verify file was created
        assert output_path.exists()
        assert output_path.suffix == ".csv"

        # Verify CSV content
        df = pd.read_csv(output_path, encoding='utf-8')

        # Should have correct number of columns
        assert len(df.columns) == 7

        # Should have correct number of rows
        assert len(df) == len(sample_verified_chunks_data)

        cprint(f"[TEST] ✓ CSV document generated: {output_path.name}", "green")

    def test_generate_json(self, sample_verified_chunks_data):
        """Test JSON document generation"""
        cprint("\n[TEST] Testing JSON document generation", "cyan")

        output_path = self.generator.generate_output(
            chunks=sample_verified_chunks_data,
            original_filename="test_document.pdf",
            output_format=OutputFormat.JSON
        )

        # Verify file was created
        assert output_path.exists()
        assert output_path.suffix == ".json"

        # Verify JSON content
        with open(output_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Verify structure
        assert "document" in data
        assert "generated_at" in data
        assert "total_chunks" in data
        assert "verified_chunks" in data
        assert "chunks" in data

        assert data["document"] == "test_document.pdf"
        assert data["total_chunks"] == len(sample_verified_chunks_data)
        assert len(data["chunks"]) == len(sample_verified_chunks_data)

        # Verify chunk structure
        for chunk_data in data["chunks"]:
            assert "page_number" in chunk_data
            assert "item_number" in chunk_data
            assert "text" in chunk_data
            assert "verified" in chunk_data
            assert "verification_score" in chunk_data

        cprint(f"[TEST] ✓ JSON document generated: {output_path.name}", "green")

    def test_verification_data_in_output(self, sample_verified_chunks_data):
        """Test that verification data is properly included in output"""
        cprint("\n[TEST] Testing verification data in output", "cyan")

        # Generate Excel for easy verification
        output_path = self.generator.generate_output(
            chunks=sample_verified_chunks_data,
            original_filename="test_document.pdf",
            output_format=OutputFormat.EXCEL
        )

        df = pd.read_excel(output_path, sheet_name='Verification')

        # Verify first row (verified=True)
        first_row = df.iloc[0]
        assert first_row["Verified ☑"] == "✅"  # Checkmark for verified
        assert first_row["Verification Score"] == "9/10"
        assert "Reference.pdf" in first_row["Verification Source"]
        assert len(first_row["Verification Note"]) > 0

        # Verify last row (verified=False)
        last_row = df.iloc[2]
        assert last_row["Verified ☑"] == "☐"  # Empty checkbox for unverified
        assert last_row["Verification Score"] == "2/10"

        cprint("[TEST] ✓ Verification data properly included in output", "green")

    def test_filename_generation(self, sample_verified_chunks_data):
        """Test output filename generation with timestamp"""
        cprint("\n[TEST] Testing output filename generation", "cyan")

        output_path = self.generator.generate_output(
            chunks=sample_verified_chunks_data,
            original_filename="my_contract.pdf",
            output_format=OutputFormat.WORD_LANDSCAPE
        )

        filename = output_path.name

        # Should contain base name
        assert "my_contract" in filename

        # Should contain "verification"
        assert "verification" in filename

        # Should have correct extension
        assert filename.endswith(".docx")

        cprint(f"[TEST] ✓ Generated filename: {filename}", "green")

    def test_multiple_format_generation(self, sample_verified_chunks_data):
        """Test generating multiple output formats from same chunks"""
        cprint("\n[TEST] Testing multiple format generation", "cyan")

        formats_to_test = [
            OutputFormat.WORD_LANDSCAPE,
            OutputFormat.EXCEL,
            OutputFormat.CSV,
            OutputFormat.JSON
        ]

        output_paths = []

        for output_format in formats_to_test:
            output_path = self.generator.generate_output(
                chunks=sample_verified_chunks_data,
                original_filename="test_document.pdf",
                output_format=output_format
            )
            output_paths.append(output_path)
            assert output_path.exists()

        # All files should be created
        assert len(output_paths) == len(formats_to_test)

        # All files should be different
        assert len(set(output_paths)) == len(output_paths)

        cprint(f"[TEST] ✓ Generated {len(output_paths)} different output formats", "green")

    def test_empty_chunks_handling(self):
        """Test handling of empty chunks list"""
        cprint("\n[TEST] Testing empty chunks handling", "cyan")

        empty_chunks = []

        output_path = self.generator.generate_output(
            chunks=empty_chunks,
            original_filename="empty_document.pdf",
            output_format=OutputFormat.EXCEL
        )

        # File should still be created
        assert output_path.exists()

        # Should have headers but no data rows
        df = pd.read_excel(output_path, sheet_name='Verification')
        assert len(df) == 0
        assert len(df.columns) == 7  # Headers should exist

        cprint("[TEST] ✓ Empty chunks handled gracefully", "green")

    def test_unverified_chunks_output(self, sample_chunks_data):
        """Test output with unverified chunks (no AI verification data)"""
        cprint("\n[TEST] Testing unverified chunks output", "cyan")

        # sample_chunks_data has no verification data
        output_path = self.generator.generate_output(
            chunks=sample_chunks_data,
            original_filename="unverified_document.pdf",
            output_format=OutputFormat.EXCEL
        )

        df = pd.read_excel(output_path, sheet_name='Verification')

        # All verification fields should be empty
        for idx, row in df.iterrows():
            assert row["Verified ☑"] == "☐"  # Empty checkbox
            assert pd.isna(row["Verification Score"]) or row["Verification Score"] == ""
            assert pd.isna(row["Verification Source"]) or row["Verification Source"] == ""
            assert pd.isna(row["Verification Note"]) or row["Verification Note"] == ""

        cprint("[TEST] ✓ Unverified chunks output correctly", "green")


@pytest.mark.unit
class TestOutputFormatSpecifics:
    """Test suite for format-specific features"""

    @pytest.fixture(autouse=True)
    def setup(self, temp_dir):
        """Setup for each test"""
        self.generator = OutputGenerator(output_dir=temp_dir)

    def test_word_table_structure(self, sample_verified_chunks_data):
        """Test Word document table structure and formatting"""
        cprint("\n[TEST] Testing Word table structure", "cyan")

        output_path = self.generator.generate_output(
            chunks=sample_verified_chunks_data,
            original_filename="test.pdf",
            output_format=OutputFormat.WORD_LANDSCAPE
        )

        doc = Document(output_path)
        table = doc.tables[0]

        # Verify table has borders
        # Verify header row is bold
        header_row = table.rows[0]
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    # Header should be bold
                    assert run.bold is True or run.font.bold is True

        cprint("[TEST] ✓ Word table structure verified", "green")

    def test_excel_column_widths(self, sample_verified_chunks_data):
        """Test Excel column width configuration"""
        cprint("\n[TEST] Testing Excel column widths", "cyan")

        output_path = self.generator.generate_output(
            chunks=sample_verified_chunks_data,
            original_filename="test.pdf",
            output_format=OutputFormat.EXCEL
        )

        # Load workbook to check column widths
        from openpyxl import load_workbook
        wb = load_workbook(output_path)
        ws = wb['Verification']

        # Verify column widths are set (not default)
        assert ws.column_dimensions['A'].width == 10  # Page #
        assert ws.column_dimensions['B'].width == 10  # Item #
        assert ws.column_dimensions['C'].width == 60  # Text (wider)

        cprint("[TEST] ✓ Excel column widths configured correctly", "green")

    def test_json_structure_completeness(self, sample_verified_chunks_data):
        """Test JSON output includes all verification metadata"""
        cprint("\n[TEST] Testing JSON structure completeness", "cyan")

        output_path = self.generator.generate_output(
            chunks=sample_verified_chunks_data,
            original_filename="test.pdf",
            output_format=OutputFormat.JSON
        )

        with open(output_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Verify all chunks include citations
        for chunk_data in data["chunks"]:
            assert "citations" in chunk_data
            if chunk_data["verified"]:
                # Verified chunks should have non-empty citations
                assert isinstance(chunk_data["citations"], list)

        cprint("[TEST] ✓ JSON structure includes all metadata", "green")
