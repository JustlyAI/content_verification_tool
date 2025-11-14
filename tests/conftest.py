"""
Pytest configuration and shared fixtures for Content Verification Tool tests
"""
import os
import pytest
import tempfile
from pathlib import Path
from typing import Dict, Any
from dotenv import load_dotenv
from termcolor import cprint

from google import genai
from google.genai import types

# Load environment variables
load_dotenv()


@pytest.fixture(scope="session")
def test_data_dir() -> Path:
    """
    Returns the path to the test data directory
    """
    return Path(__file__).parent / "data"


@pytest.fixture(scope="session")
def sample_pdf_path(test_data_dir: Path) -> Path:
    """
    Returns path to sample PDF file for testing
    """
    return test_data_dir / "sample_document.pdf"


@pytest.fixture(scope="session")
def sample_docx_path(test_data_dir: Path) -> Path:
    """
    Returns path to sample DOCX file for testing
    """
    return test_data_dir / "sample_document.docx"


@pytest.fixture(scope="session")
def sample_reference_pdf_path(test_data_dir: Path) -> Path:
    """
    Returns path to sample reference document for testing
    """
    return test_data_dir / "reference_document.pdf"


@pytest.fixture(scope="session")
def gemini_api_key() -> str:
    """
    Returns Gemini API key from environment
    Skips tests if API key is not available
    """
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        pytest.skip("GEMINI_API_KEY not found in environment variables")
    return api_key


@pytest.fixture(scope="session")
def gemini_client(gemini_api_key: str):
    """
    Returns initialized Gemini client for testing
    """
    client = genai.Client(api_key=gemini_api_key)
    cprint("✓ Gemini test client initialized", "green")
    return client


@pytest.fixture
def temp_dir():
    """
    Creates a temporary directory for test outputs
    Cleans up after test
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        yield Path(tmp_dir)


@pytest.fixture
def sample_document_content() -> bytes:
    """
    Returns sample PDF document content in bytes
    Uses a minimal PDF for fast testing
    """
    # Minimal valid PDF content (blank 1-page PDF)
    pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/Resources <<
/Font <<
/F1 <<
/Type /Font
/Subtype /Type1
/BaseFont /Helvetica
>>
>>
>>
/MediaBox [0 0 612 792]
/Contents 4 0 R
>>
endobj
4 0 obj
<<
/Length 44
>>
stream
BT
/F1 12 Tf
100 700 Td
(Test Document) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000317 00000 n
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
410
%%EOF
"""
    return pdf_content


@pytest.fixture
def sample_docx_content() -> bytes:
    """
    Returns sample DOCX document content in bytes
    Creates a minimal valid DOCX for testing
    """
    from docx import Document
    import io

    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph for verification.')
    doc.add_paragraph('This is another sentence for testing.')
    doc.add_paragraph('Contract terms: The party agrees to all conditions.')

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


@pytest.fixture
def sample_chunks_data() -> list:
    """
    Returns sample document chunks for testing
    """
    from backend.app.models import DocumentChunk

    return [
        DocumentChunk(
            page_number=1,
            item_number="1",
            text="This is the first paragraph of the document.",
            is_overlap=False
        ),
        DocumentChunk(
            page_number=1,
            item_number="2",
            text="This is the second paragraph on page one.",
            is_overlap=False
        ),
        DocumentChunk(
            page_number=2,
            item_number="1",
            text="This paragraph starts on page two.",
            is_overlap=False
        ),
    ]


@pytest.fixture
def sample_verified_chunks_data() -> list:
    """
    Returns sample verified chunks for testing output generation
    """
    from backend.app.models import DocumentChunk

    return [
        DocumentChunk(
            page_number=1,
            item_number="1",
            text="The contract was signed on January 15, 2024.",
            is_overlap=False,
            verified=True,
            verification_score=9,
            verification_source="Reference.pdf, Section 2.1",
            verification_note="Verified against reference document",
            citations=[
                {
                    "title": "Reference.pdf",
                    "excerpt": "The contract was signed on January 15, 2024."
                }
            ]
        ),
        DocumentChunk(
            page_number=1,
            item_number="2",
            text="The party agrees to all terms and conditions.",
            is_overlap=False,
            verified=True,
            verification_score=8,
            verification_source="Reference.pdf, Section 3.2",
            verification_note="Confirmed in terms section",
            citations=[
                {
                    "title": "Reference.pdf",
                    "excerpt": "all terms and conditions"
                }
            ]
        ),
        DocumentChunk(
            page_number=1,
            item_number="3",
            text="This statement has no supporting evidence.",
            is_overlap=False,
            verified=False,
            verification_score=2,
            verification_source="No source found",
            verification_note="Could not find supporting documentation",
            citations=[]
        ),
    ]


@pytest.fixture
def mock_file_search_store(gemini_client) -> str:
    """
    Creates a temporary File Search store for testing
    Returns store name
    Cleans up after test
    """
    # Create test store
    store = gemini_client.file_search_stores.create(
        config={'display_name': f'Test Store {os.urandom(4).hex()}'}
    )

    cprint(f"✓ Created test File Search store: {store.name}", "green")

    yield store.name

    # Cleanup: Delete store after test
    try:
        gemini_client.file_search_stores.delete(name=store.name)
        cprint(f"✓ Deleted test File Search store: {store.name}", "green")
    except Exception as e:
        cprint(f"⚠️  Could not delete test store: {e}", "yellow")


@pytest.fixture
def sample_case_context() -> str:
    """
    Returns sample case context for testing
    """
    return """
    This is a contract verification case. The document under review is a service
    agreement between Company A and Company B dated January 2024. We need to verify
    that all claims made in the document are supported by the reference materials.
    """


@pytest.fixture
def sample_metadata() -> Dict[str, Any]:
    """
    Returns sample metadata for testing
    """
    from backend.app.models import DocumentMetadata
    from datetime import datetime

    return DocumentMetadata(
        document_id="test_doc_123",
        filename="test_reference.pdf",
        summary="This is a test reference document containing contract information.",
        contextualization="This document provides evidence for contract verification.",
        document_type="contract",
        keywords=["contract", "agreement", "terms", "conditions", "signature"],
        generated_at=datetime.now()
    )


@pytest.fixture(scope="session", autouse=True)
def setup_test_environment():
    """
    Setup test environment before running tests
    """
    cprint("\n" + "=" * 80, "cyan")
    cprint("🧪 Setting up Content Verification Tool Test Suite", "green", attrs=["bold"])
    cprint("=" * 80, "cyan")

    # Verify environment
    if not os.getenv("GEMINI_API_KEY"):
        cprint("⚠️  WARNING: GEMINI_API_KEY not set - some tests will be skipped", "yellow")
    else:
        cprint("✓ GEMINI_API_KEY found", "green")

    cprint("=" * 80 + "\n", "cyan")

    yield

    cprint("\n" + "=" * 80, "cyan")
    cprint("🧪 Test Suite Complete", "green", attrs=["bold"])
    cprint("=" * 80 + "\n", "cyan")


# Pytest configuration
def pytest_configure(config):
    """
    Configure pytest with custom markers
    """
    config.addinivalue_line(
        "markers", "integration: marks tests as integration tests (may be slow)"
    )
    config.addinivalue_line(
        "markers", "unit: marks tests as unit tests (fast)"
    )
    config.addinivalue_line(
        "markers", "gemini: marks tests that require Gemini API access"
    )
    config.addinivalue_line(
        "markers", "slow: marks tests as slow (deselect with '-m \"not slow\"')"
    )
