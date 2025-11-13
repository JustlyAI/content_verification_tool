"""
Output generation module for Word and Excel/CSV formats
"""
from pathlib import Path
from typing import List
from datetime import datetime
import json
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from termcolor import cprint

from app.models import DocumentChunk, OutputFormat


# Output directory
OUTPUT_DIR = Path("/tmp/output")
OUTPUT_DIR.mkdir(exist_ok=True)


class OutputGenerator:
    """Handles generation of verification documents in various formats"""

    def __init__(self, output_dir: Path = OUTPUT_DIR):
        self.output_dir = output_dir
        self.output_dir.mkdir(exist_ok=True)
        cprint(f"[OUTPUT] Initialized output directory: {self.output_dir}", "cyan")

    def _generate_filename(self, original_filename: str, output_format: OutputFormat) -> str:
        """
        Generate output filename with timestamp

        Args:
            original_filename: Original document filename
            output_format: Output format

        Returns:
            Generated filename
        """
        base_name = Path(original_filename).stem
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        if output_format in [OutputFormat.WORD_LANDSCAPE, OutputFormat.WORD_PORTRAIT]:
            ext = "docx"
        elif output_format == OutputFormat.EXCEL:
            ext = "xlsx"
        elif output_format == OutputFormat.CSV:
            ext = "csv"
        elif output_format == OutputFormat.JSON:
            ext = "json"
        else:
            ext = "txt"

        return f"{base_name}_verification_{timestamp}.{ext}"

    def _set_cell_border(self, cell, **kwargs):
        """
        Set cell border in Word table

        Args:
            cell: Table cell object
            **kwargs: Border properties
        """
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        # Create border element
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '4')
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), '000000')
            tcBorders.append(edge_element)

        tcPr.append(tcBorders)

    def generate_word_document(
        self,
        chunks: List[DocumentChunk],
        original_filename: str,
        landscape: bool = True
    ) -> Path:
        """
        Generate Word document with verification table

        Args:
            chunks: List of document chunks
            original_filename: Original document filename
            landscape: True for landscape, False for portrait orientation

        Returns:
            Path to generated document
        """
        orientation = "landscape" if landscape else "portrait"
        cprint(f"[OUTPUT] Generating Word document ({orientation})...", "cyan")

        # Create document
        doc = Document()

        # Set page orientation
        section = doc.sections[0]
        if landscape:
            # Landscape: swap width and height
            section.page_width = Inches(11)
            section.page_height = Inches(8.5)
        else:
            # Portrait (default)
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

        # Set margins
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        # Add title
        title = doc.add_paragraph()
        title_run = title.add_run(f"Verification Document: {Path(original_filename).stem}")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

        # Create table with 6 columns
        # Add 1 for header row
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'

        # Set column widths based on orientation
        if landscape:
            # Landscape: more space for text and notes
            col_widths = [Inches(0.7), Inches(0.7), Inches(4.0), Inches(0.8), Inches(1.8), Inches(1.8)]
        else:
            # Portrait: compressed layout
            col_widths = [Inches(0.6), Inches(0.6), Inches(2.5), Inches(0.7), Inches(1.5), Inches(1.5)]

        for idx, width in enumerate(col_widths):
            for cell in table.columns[idx].cells:
                cell.width = width

        # Set header row
        header_cells = table.rows[0].cells
        headers = ["Page #", "Item #", "Text", "Verified ☑", "Verification Source", "Verification Note"]

        for idx, header_text in enumerate(headers):
            cell = header_cells[idx]
            cell.text = header_text
            # Make header bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add data rows
        for chunk in chunks:
            row_cells = table.add_row().cells

            # Page number
            row_cells[0].text = str(chunk.page_number)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Item number
            row_cells[1].text = str(chunk.item_number)
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Text
            row_cells[2].text = chunk.text
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Verified checkbox - show AI result if available
            if chunk.verified is not None:
                row_cells[3].text = "✅" if chunk.verified else "☐"
            else:
                row_cells[3].text = "☐"
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Verification Source - populate from AI if available
            if chunk.verification_source:
                row_cells[4].text = chunk.verification_source
            else:
                row_cells[4].text = ""

            # Verification Note - include AI reasoning and confidence score
            if chunk.verification_note:
                note_text = chunk.verification_note
                if chunk.verification_score:
                    note_text += f" (Confidence: {chunk.verification_score}/10)"
                row_cells[5].text = note_text
            else:
                row_cells[5].text = ""

            # Set font size for all cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        # Generate filename and save
        output_format = OutputFormat.WORD_LANDSCAPE if landscape else OutputFormat.WORD_PORTRAIT
        filename = self._generate_filename(original_filename, output_format)
        output_path = self.output_dir / filename

        doc.save(output_path)
        cprint(f"[OUTPUT] Word document saved: {output_path}", "green")

        return output_path

    def generate_excel_csv(
        self,
        chunks: List[DocumentChunk],
        original_filename: str,
        as_excel: bool = True
    ) -> Path:
        """
        Generate Excel or CSV file with verification table

        Args:
            chunks: List of document chunks
            original_filename: Original document filename
            as_excel: True for Excel, False for CSV

        Returns:
            Path to generated file
        """
        file_type = "Excel" if as_excel else "CSV"
        cprint(f"[OUTPUT] Generating {file_type} file...", "cyan")

        # Create DataFrame with AI verification results
        verified_values = []
        source_values = []
        note_values = []

        for chunk in chunks:
            # Verified column - show checkmark if AI verified
            if chunk.verified is not None:
                verified_values.append("✅" if chunk.verified else "☐")
            else:
                verified_values.append("☐")

            # Source column - populate from AI if available
            source_values.append(chunk.verification_source if chunk.verification_source else "")

            # Note column - include AI reasoning and confidence score
            if chunk.verification_note:
                note_text = chunk.verification_note
                if chunk.verification_score:
                    note_text += f" (Confidence: {chunk.verification_score}/10)"
                note_values.append(note_text)
            else:
                note_values.append("")

        data = {
            "Page #": [chunk.page_number for chunk in chunks],
            "Item #": [chunk.item_number for chunk in chunks],
            "Text": [chunk.text for chunk in chunks],
            "Verified ☑": verified_values,
            "Verification Source": source_values,
            "Verification Note": note_values
        }

        df = pd.DataFrame(data)

        # Generate filename and save
        output_format = OutputFormat.EXCEL if as_excel else OutputFormat.CSV
        filename = self._generate_filename(original_filename, output_format)
        output_path = self.output_dir / filename

        if as_excel:
            # Save as Excel with basic formatting
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Verification')

                # Get worksheet
                worksheet = writer.sheets['Verification']

                # Set column widths
                column_widths = {
                    'A': 10,  # Page #
                    'B': 10,  # Item #
                    'C': 60,  # Text
                    'D': 12,  # Verified
                    'E': 25,  # Verification Source
                    'F': 25   # Verification Note
                }

                for col, width in column_widths.items():
                    worksheet.column_dimensions[col].width = width

            cprint(f"[OUTPUT] Excel file saved: {output_path}", "green")
        else:
            # Save as CSV
            df.to_csv(output_path, index=False, encoding='utf-8')
            cprint(f"[OUTPUT] CSV file saved: {output_path}", "green")

        return output_path

    def generate_json(
        self,
        chunks: List[DocumentChunk],
        original_filename: str
    ) -> Path:
        """
        Generate JSON file with full verification metadata

        Args:
            chunks: List of document chunks
            original_filename: Original document filename

        Returns:
            Path to generated JSON file
        """
        cprint(f"[OUTPUT] Generating JSON file...", "cyan")

        # Convert chunks to dict format
        chunks_data = []
        for chunk in chunks:
            chunk_dict = {
                "page_number": chunk.page_number,
                "item_number": chunk.item_number,
                "text": chunk.text,
                "is_overlap": chunk.is_overlap,
                "verified": chunk.verified,
                "verification_score": chunk.verification_score,
                "verification_source": chunk.verification_source,
                "verification_note": chunk.verification_note,
                "citations": chunk.citations
            }
            chunks_data.append(chunk_dict)

        # Create output structure
        output_data = {
            "document": original_filename,
            "generated_at": datetime.now().isoformat(),
            "total_chunks": len(chunks),
            "verified_chunks": sum(1 for c in chunks if c.verified),
            "chunks": chunks_data
        }

        # Generate filename and save
        filename = self._generate_filename(original_filename, OutputFormat.JSON)
        output_path = self.output_dir / filename

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False)

        cprint(f"[OUTPUT] JSON file saved: {output_path}", "green")
        return output_path

    def generate_output(
        self,
        chunks: List[DocumentChunk],
        original_filename: str,
        output_format: OutputFormat
    ) -> Path:
        """
        Generate output in specified format

        Args:
            chunks: List of document chunks
            original_filename: Original document filename
            output_format: Desired output format

        Returns:
            Path to generated file
        """
        cprint(f"[OUTPUT] Generating output in {output_format.value} format...", "cyan")

        if output_format == OutputFormat.WORD_LANDSCAPE:
            return self.generate_word_document(chunks, original_filename, landscape=True)
        elif output_format == OutputFormat.WORD_PORTRAIT:
            return self.generate_word_document(chunks, original_filename, landscape=False)
        elif output_format == OutputFormat.EXCEL:
            return self.generate_excel_csv(chunks, original_filename, as_excel=True)
        elif output_format == OutputFormat.CSV:
            return self.generate_excel_csv(chunks, original_filename, as_excel=False)
        elif output_format == OutputFormat.JSON:
            return self.generate_json(chunks, original_filename)
        else:
            raise ValueError(f"Unknown output format: {output_format}")


# Global output generator instance
output_generator = OutputGenerator()
